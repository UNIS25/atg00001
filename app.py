import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

st.set_page_config(page_title="United Nations Information Services Digitalization Project", layout="wide")

UN_LOGO = "https://logowik.com/content/uploads/images/un-united-nations3511.logowik.com.webp"

if "page" not in st.session_state:
    st.session_state.page = "Home"

st.markdown("""
    <style>
        body, .stApp {
            background-color: #009EDB !important;
        }
        .header {
            display: flex;
            align-items: center;
            justify-content: start;
            background-color: white !important;
            padding: 20px 50px;
            border-bottom: 5px solid #0072C6;
            box-shadow: 0px 4px 8px rgba(0, 0, 0, 0.1);
            color: black !important;
        }
        .header img {
            width: 240px;
            margin-right: 20px;
        }
        .header-content {
            display: flex;
            flex-direction: column;
            justify-content: center;
            color: black !important;
        }
        .header-text {
            font-size: 28px;
            font-weight: bold;
            text-align: left;
        }
        .tagline {
            font-size: 18px;
            font-weight: bold;
            text-align: left;
        }
        .developer {
            font-size: 12px;
            margin-top: 5px;
        }
        .container {
            text-align: center;
            padding: 50px;
            background-color: #009EDB !important;
            color: yellow !important;
            border-radius: 10px;
            width: 80%;
            margin: auto;
        }
        h2, h3, h4, h5, h6, p, .stMarkdown a {
            color: yellow !important;
        }
        .st-emotion-cache-10trblm, .st-emotion-cache-1v0mbdj p, .stMarkdown p {
            color: black !important;
        }
    </style>
""", unsafe_allow_html=True)

st.markdown(f"""
    <div class="header">
        <img src="{UN_LOGO}" alt="UN Logo">
        <div class="header-content">
            <div class="header-text">United Nations Information Services Digitalization Project</div>
            <div class="tagline">Transforming traditional documentation into a digital future</div>
            <div class="developer">Developed by ATG</div>
        </div>
    </div>
""", unsafe_allow_html=True)

selected_page = st.radio("", ["Home", "Weekly Report", "Contact"], horizontal=True)
st.session_state.page = selected_page

if st.session_state.page == "Home":
    st.markdown('<div class="container">', unsafe_allow_html=True)
    st.markdown("""
        <h2>Welcome to the Digital Future</h2>
        <p>Explore the power of digital transformation in streamlining UN documentation.</p>
        <p><strong>Discover UN BRAIN</strong> — our new initiative to build secure, sovereign, AI-powered tools for the United Nations, developed with digital independence in mind.</p>
        <p>Click below to read more about this innovation:</p>
        <p><a href="https://www.linkedin.com/posts/adrian-veluppillai-312550199_aiforgood-untech-digitalsovereignty-activity-7325115202822180864-fdez?utm_source=share&utm_medium=member_desktop&rcm=ACoAAC6IYr4B2nyUuSXe8V1doalv1dOuh5TOe6s" target="_blank">🔗 Click here for more information</a></p>
        <img src="https://media.licdn.com/dms/image/v2/D4E22AQEYBf9eJQRjww/feedshare-shrink_2048_1536/B4EZagJdJ.HIAw-/0/1746443556142?e=1749686400&v=beta&t=50wrYQoEe6bpxZxBZU3_YYT5oMhdfFeuAMxs44aqgJM" width="80%">
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.page == "Weekly Report":
    st.markdown("<div class='container'><h2>Weekly Report</h2><p>Upload your files and analyze engagement data.</p></div>", unsafe_allow_html=True)

    def read_file(uploaded_file, platform=None):
        if uploaded_file.name.endswith(".csv"):
            return pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith((".xls", ".xlsx")):
            if platform == "LinkedIn":
                return pd.read_excel(uploaded_file, sheet_name="All posts", skiprows=1)
            return pd.read_excel(uploaded_file)
        else:
            st.error("Unsupported file format. Upload .csv, .xls, or .xlsx files.")
            return None

    def map_columns(df, column_map):
        new_mapping = {}
        for key, value in column_map.items():
            possible_columns = key.split("|")
            for col in possible_columns:
                if col in df.columns:
                    new_mapping[col] = value
                    break
        return new_mapping

    def process_platform(df, platform, cols_map, engagement_cols, start_date, end_date):
        cols_map = map_columns(df, cols_map)
        df_selected = df[list(cols_map.keys())].rename(columns=cols_map)
        df_selected['Date'] = pd.to_datetime(df_selected['Date'], errors='coerce')
        start_date = pd.to_datetime(start_date)
        end_date = pd.to_datetime(end_date)
        df_selected = df_selected[(df_selected['Date'] >= start_date) & (df_selected['Date'] <= end_date)]
        df_selected['Engagements'] = df_selected[engagement_cols].sum(axis=1)
        df_selected = df_selected.sort_values(by='Engagements', ascending=False).head(3)
        df_selected.insert(0, 'Platform', platform)
        df_selected.insert(1, 'Rank', [1, 2, 3])
        df_selected['Date'] = df_selected['Date'].dt.strftime('%m/%d/%Y')
        return df_selected[["Platform", "Rank", "Date", "Post text", "Link", "Impressions", "Engagements", "Reactions", "Comments", "Shares"]]

    st.sidebar.header("Select Date Range")
    start_date = st.sidebar.date_input("Start date")
    end_date = st.sidebar.date_input("End date")

    processing_functions = {
        "X English": lambda df: process_platform(df.rename(columns={"Link": "Post Link"}), "X English", {
            "Date": "Date", "Post text": "Post text", "Post Link": "Link", "Impressions": "Impressions",
            "Engagements": "Engagements", "Likes": "Reactions", "Replies": "Comments", "Reposts": "Shares"
        }, ["Engagements"], start_date, end_date),

        "X French": lambda df: process_platform(df.rename(columns={"Lien": "Lien du post"}), "X French", {
            "Date": "Date", "Texte du post": "Post text", "Lien du post": "Link", "Impressions": "Impressions",
            "Engagements": "Engagements", "J'aime": "Reactions", "Réponses": "Comments", "Reposts": "Shares"
        }, ["Engagements"], start_date, end_date),

        "Facebook": lambda df: process_platform(df, "Facebook", {
            "Publish time|Heure de publication": "Date", "Title|Titre": "Post text", "Permalink|Permalien": "Link", 
            "Reach|Couverture": "Impressions", "Reactions|Réactions": "Reactions", "Comments|Commentaires": "Comments", 
            "Shares|Partages": "Shares"
        }, ["Reactions", "Comments", "Shares"], start_date, end_date),

        "Instagram": lambda df: process_platform(df, "Instagram", {
            "Publish time|Heure de publication": "Date", "Description": "Post text", "Permalink|Permalien": "Link", 
            "Reach|Couverture": "Impressions", "Likes|Mentions J’aime": "Reactions", "Shares|Partages": "Shares", 
            "Follows|Followers en plus": "Follows", "Comments|Commentaires": "Comments", "Saves|Enregistrements": "Saves"
        }, ["Reactions", "Shares", "Follows", "Comments", "Saves"], start_date, end_date),

        "LinkedIn": lambda df: process_platform(df, "LinkedIn", {
            "Created date": "Date", "Post title": "Post text", "Post link": "Link", "Impressions": "Impressions",
            "Clicks": "Clicks", "Likes": "Reactions", "Comments": "Comments", "Reposts": "Shares", "Follows": "Follows"
        }, ["Clicks", "Reactions", "Comments", "Shares", "Follows"], start_date, end_date)
    }

    platform_dfs = {}
    for platform in processing_functions:
        uploaded_file = st.file_uploader(f'Upload file for {platform}', type=['csv', 'xls', 'xlsx'], key=platform)
        if uploaded_file:
            df = read_file(uploaded_file, platform=platform)
            if df is not None:
                platform_df = processing_functions[platform](df)
                platform_dfs[platform] = platform_df
                st.write(f"Top 3 Posts for {platform}")
                st.dataframe(platform_df)
                st.markdown('---')

    if platform_dfs:
        combined_df = pd.concat(platform_dfs.values(), ignore_index=True)
        st.write("### Combined Top Posts Report")
        st.dataframe(combined_df)

        combined_df.to_excel("final_social_report.xlsx", index=False)
        with open("final_social_report.xlsx", "rb") as f:
            st.download_button("Download Excel Report", f, "Social_Media_Report.xlsx")

        if st.button("Open Tabs"):
            st.markdown("### Click below to open each post:")
            for idx, row in combined_df.iterrows():
                platform = row["Platform"]
                rank = row["Rank"]
                link = row["Link"]
                st.markdown(f"- [{platform} Rank {rank} Post]({link})", unsafe_allow_html=True)

        if st.button("Generate Word Report"):
            doc = Document()
            doc.add_heading('Weekly Social Media Performance Report', level=1)
            doc.add_heading('1. Summary', level=2)
            doc.add_paragraph("This report highlights the top-performing posts across social media platforms. The analysis includes engagement metrics, impressions, and trends.")
            combined_df["Date"] = pd.to_datetime(combined_df["Date"], errors='coerce').dt.strftime("%d/%m/%Y")

            for platform in combined_df["Platform"].unique():
                doc.add_heading(f'Top Posts - {platform}', level=2)
                platform_df = combined_df[combined_df["Platform"] == platform].nlargest(3, 'Engagements')
                for idx, row in platform_df.iterrows():
                    doc.add_paragraph(f"{idx+1}. {row['Post text']}", style="List Bullet")
                    doc.add_paragraph(f"- Date: {row['Date']}\n")
                    doc.add_paragraph(f"- Impressions: {row['Impressions']}\n")
                    doc.add_paragraph(f"- Engagements: {row['Engagements']}\n")
                    doc.add_paragraph(f"- Reactions: {row['Reactions']} | Comments: {row['Comments']} | Shares: {row['Shares']}\n")
                    doc.add_paragraph(f"- View Post: {row['Link']}\n")

            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)
            st.download_button("Download Word Report", word_buffer, file_name="Weekly_Social_Media_Report.docx")

if st.session_state.page == "Contact":
    st.markdown('<div class="container"><h2>Contact Us</h2><p>For inquiries, reach out to:</p><p><strong>Email:</strong> <a href="mailto:adrian.vasuveluppillai@un.org">adrian.vasuveluppillai@un.org</a></p></div>', unsafe_allow_html=True)
